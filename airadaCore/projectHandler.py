import logging 

import docx
from docx.table import (
    Table, _Row
)
from python_docx_replace import docx_replace

from airadaCore import utils

from typing import Any
from itertools import chain
from airadaCore.airadaTypes import (
    Path, jsonData
)

OUTPUT_PATH = "./output/out.docx"
PROJECT_PAPER_TEMPLATE_PATH = './template/project-paper.docx'

STEP_TI = 1
BUDGET_TI = 2
OKR_TI = 3
ACTIVITY_CREDIT_TI = 4
SKILL_CREDIT_TI = 5

BLANK_CHECKBOX = "☐"
CHECKED_CHECKBOX = "☒"

logger: logging.Logger = logging.getLogger("airadaCore/projectHandler")


def handle_credit_section(document: docx.Document, data: dict[str, Any]) -> None:
    # Filter only row with templated blank checkbox ☐
    def get_clickables(table: Table) -> tuple[_Row]:
        return tuple(row for row in table.rows if row.cells[0].text == BLANK_CHECKBOX)
    
    
    # Replace blank checkbox ☐ with ☒ if the value is true for each row
    def check_credit_table(answers: list[dict[str, bool]], checkables: list[_Row]) -> tuple[_Row]:
        for answer, checkable in zip(answers, checkables):
            # organizer and participant are in cell 0 and 1 of each row
            for i, key in enumerate(("organizer", "participant")):
                if answer[key]:
                    checkable.cells[i].paragraphs[0].text = CHECKED_CHECKBOX
    
    
    activity_checkables: tuple[_Row] = get_clickables(document.tables[ACTIVITY_CREDIT_TI])
    check_credit_table(data["activityCredits"], activity_checkables)
    
    skill_checkables: tuple[_Row] = get_clickables(document.tables[SKILL_CREDIT_TI])
    check_credit_table(data["skillCredits"], skill_checkables)
    
    # get text of most expected skill section based on the text in the able → allow for easy adjustment
    most_expected_text: str = skill_checkables[data["mostExpectedSkill"]].cells[2].text.replace("\xa0\n", "")
    docx_replace(document, mostExpectedSkill=most_expected_text)


def add_tabel_placeholders(n_row: int, table: Table, swap_last: bool = False) -> None:
    # Becasue template row was already in place, so only n-1 new rows needed.
    # rows[0] = header; rows[1] = template
    for _ in range(n_row - 1):
        # create new row and then duplicate content from template
        new_row: _Row = table.add_row()

        for cell, template_cell in zip(new_row.cells, table.rows[1].cells):
            cell.paragraphs[0].text = template_cell.paragraphs[0].text
            cell.paragraphs[0].style = template_cell.paragraphs[0].style

        # Swap to before last position. Primary used in budget table
        if (swap_last): 
            table.rows[-3]._tr.addnext(new_row._tr)

    # replace each # with correct running number
    # [1:] to skip header to preserve style
    # also skip final row [1:-1] if swap_last == True
    rows = table.rows[1:-1] if swap_last else table.rows[1:]
    
    for i, row in enumerate(rows) :
        for cell in row.cells:
            cell.paragraphs[0].text = cell.paragraphs[0].text.replace("#", str(i))


def replace_paragraph_placeholders(document: docx.Document, topics: dict[str, list|tuple]) -> None:
    # k[:-1] to strip the final -s
    def get_texts(topic, texts) -> tuple[str, str]:
        return tuple((f"{topic[:-1]} {i}", text) for i, text in enumerate(texts))
    
    
    # flatten the lists and assign to map
    _ = dict(chain.from_iterable(get_texts(k, topics[k]) for k in topics))
    docx_replace(document, **_)
    

def replace_inline_placeholders(document: docx.Document, data: dict[str, str|int]) -> None:
    # Non-numerated placeholders is simply replaced with text
    # i.e., topics that are inline text or could not have multiple paragraphs
    # coresponded with non-iterable values (that is not str) in processed_data
    docx_replace(document, **data)


def prepare_paragraphs(document: docx.Document, topics: dict[str, list|tuple]) -> None:
    # Enumerated placeholders. Placeholder index starts at 0.
    # k[:-1] to strip the final -s
    def get_sub_placeholders_text(topic: str, n: int) -> str:
        return "\n\t".join(f"${{{topic[:-1]} {i}}}" for i in range(n))
    
    
    # replace placeholder ending with -s with multiple placeholder with running number
    # e.g., from    ${ABCs} →
    #               ${ABC 0}
    #               ${ABC 1} ... and so on
    
    _: dict[str, str] = {
        topic: get_sub_placeholders_text(topic, len(topics[topic])) for topic in topics
    }
    docx_replace(document, **_)


def prepare_tables(tables: list[Table], data: dict[str, Any]) -> None:
    add_tabel_placeholders(len(data["steps"]), tables[STEP_TI])
    add_tabel_placeholders(len(data["consequencesOKRs"]), tables[OKR_TI])
    add_tabel_placeholders(len(data["budgetItems"]), tables[BUDGET_TI], swap_last=True)


def process_data(data: jsonData) -> dict[str, Any]:
    def get_student_manager_text(_: dict) -> str:
        return f"{_["name"]}\t{_["id"]}\t{_["role"]}"
    
    def get_coucil_manager_text(_: dict) -> str:
        return f"{_["name"]}\t{_["role"]}"
    
        
    # these topics do not need further modification
    # so they are simply copy from the original date
    SIMPLE_COPY_KEYS: tuple[str] = (
        "projectName", "rationales", "objectives", 
        "contentAdvisor", "technicalAdvisor",
        "nProfessor", "nStaff", "nStudent", "nExternal",
        "location",
    )
    processed_data: dict[str, str|tuple] = {k: data[k] for k in SIMPLE_COPY_KEYS}

    # these topics need some processing
    # must convert any iterable to list or tuple as len() is needed
    processed_data.update({
                            # add number of each types of participants
        "nSum":             sum(data[_] for _ in ("nProfessor", "nStaff", "nStudent", "nExternal")),
        
                            # format ISO date string to Thai
        "periodStartDate":  utils.toThaiDate(data["periodStartDate"]),
        "periodEndDate":    utils.toThaiDate(data["periodEndDate"]),

                            # format each one to correct form
        "councilManager":   get_coucil_manager_text(data["studentCouncilManager"]),
        "studentManagers":  tuple(map(get_student_manager_text, data["studentManagers"])),

                            # create a list of each items text with numbering
                            # (e.g., "1. Apple", "2. Banana")
        "budgetItems":      utils.each_to_numbered_list("item", data["budgets"]),
        
                            # create a list of each values with coresponding keys,
                            # then sum them, and create Thai text
        "budgetPrices":     (budget_prices := utils.get_each("price", data["budgets"])),
        "budgetSum":        (budget_sum := utils.sum_money_str(budget_prices)),
        "budgetSumText":    utils.get_money_thai_text(budget_sum),

        "consequencesOKRs": utils.each_to_numbered_list("OKR", data["consequences"]),
        "consequencesKPIs": utils.get_each("KPI", data["consequences"]),

        "steps":            utils.each_to_numbered_list("step", data["steps"]),
        "stepStartDates":   (step_start_dates := utils.get_each("startDate", data["steps"])),
        "stepEndDates":     (step_end_dates := utils.get_each("endDate", data["steps"])),
        
                            # get duration of each steps
        "stepPeroids":      tuple(map(utils.get_date_delta_thai_text, step_start_dates, step_end_dates)),
        "stepManagers":     utils.get_each("manager", data["steps"])
    })

    # TODO "programmes"
    return processed_data


def handle(data: jsonData) -> Path:
    logger.info(f"Output path set to {OUTPUT_PATH}")
    logger.info(f"Template path set to {PROJECT_PAPER_TEMPLATE_PATH}")
    
    # copy template to output folder
    docx.Document(PROJECT_PAPER_TEMPLATE_PATH).save(OUTPUT_PATH)

    # open copied document
    document: docx.Document = docx.Document(OUTPUT_PATH)
    
    # credit section are handled separately for simplicity
    handle_credit_section(document, data)

    # processed_data contains only texts or groups of texts which replaced template directly
    # other data that do not meet this criteria CANNOT be in processed_data (e.g., handle_credit_section)
    processed_data: dict[str, Any] = process_data(data) 
    
    inline_texts: dict[str, str|int] = {k: processed_data[k] for k in processed_data if isinstance(processed_data[k], (str, int))}
    paragraphs: dict[str, tuple|dict] = {k: processed_data[k] for k in processed_data if isinstance(processed_data[k], (list, tuple))}
    
    # create placeholders matching number of data in processed_data
    prepare_paragraphs(document, paragraphs)
    
    # similar to prepare_paragraphs. this also add rows to table.
    prepare_tables(document.tables, paragraphs)
    
    # replace each inline placeholder
    replace_inline_placeholders(document, inline_texts)
    replace_paragraph_placeholders(document, paragraphs)
    
    document.save(OUTPUT_PATH)
    
    return OUTPUT_PATH